"""
Resume PII Redactor
-------------------
Watches an incoming folder for new resumes (.txt, .docx), redacts PII
using Microsoft Presidio (runs entirely on your laptop — no cloud calls),
saves the cleaned copy to a redacted folder, and archives the original.

Folder layout (created automatically by setup_redaction.py):
    ResumePII/
        incoming/   <- drop resumes here
        redacted/   <- redacted copies appear here
        archive/    <- originals moved here after processing
        error/      <- files that could not be processed
"""

import logging
import shutil
import time
from pathlib import Path

from watchdog.events import FileSystemEventHandler
from watchdog.observers import Observer

# ---------------------------------------------------------------------------
# Configuration — edit these paths to match wherever you created the folders
# ---------------------------------------------------------------------------
BASE = Path.home() / "ResumePII"
INCOMING = BASE / "incoming"
REDACTED = BASE / "redacted"
ARCHIVE = BASE / "archive"
ERROR = BASE / "error"

# PII entity types Presidio will detect and replace.
# Full list: https://microsoft.github.io/presidio/supported_entities/
PII_ENTITIES = [
    "PERSON",
    "EMAIL_ADDRESS",
    "PHONE_NUMBER",
    "LOCATION",
    "URL",
    "DATE_TIME",
    "NRP",           # Nationality, Religious or Political group
    "MEDICAL_LICENSE",
    "US_SSN",
    "US_DRIVER_LICENSE",
    "US_PASSPORT",
    "CREDIT_CARD",
    "IBAN_CODE",
    "IP_ADDRESS",
]

# How long to wait (seconds) after a file appears before reading it.
# Prevents reading a file mid-copy.
COPY_SETTLE_SECONDS = 2

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Lazy-load Presidio so import errors surface at startup, not mid-run
# ---------------------------------------------------------------------------
try:
    from presidio_analyzer import AnalyzerEngine
    from presidio_anonymizer import AnonymizerEngine
    from presidio_anonymizer.entities import OperatorConfig
except ImportError:
    raise SystemExit(
        "\nPresidio is not installed. Run:\n"
        "  pip install -r requirements-redaction.txt\n"
        "  python -m spacy download en_core_web_lg\n"
    )

try:
    import docx  # python-docx
except ImportError:
    docx = None
    log.warning("python-docx not installed — .docx files will be sent to error/")

_analyzer = None
_anonymizer = None


def _get_engines():
    """Initialise Presidio engines once and reuse them."""
    global _analyzer, _anonymizer
    if _analyzer is None:
        log.info("Loading Presidio NLP engine (first run may take ~10 seconds)…")
        _analyzer = AnalyzerEngine()
        _anonymizer = AnonymizerEngine()
        log.info("Presidio ready.")
    return _analyzer, _anonymizer


# ---------------------------------------------------------------------------
# Core redaction helpers
# ---------------------------------------------------------------------------

def redact_text(raw_text: str) -> str:
    """Return a copy of raw_text with all detected PII replaced by labels."""
    analyzer, anonymizer = _get_engines()

    results = analyzer.analyze(
        text=raw_text,
        entities=PII_ENTITIES,
        language="en",
    )

    # Replace each PII span with a readable label like <PERSON> or <EMAIL_ADDRESS>
    operators = {
        entity: OperatorConfig("replace", {"new_value": f"<{entity}>"})
        for entity in PII_ENTITIES
    }

    anonymized = anonymizer.anonymize(
        text=raw_text,
        analyzer_results=results,
        operators=operators,
    )
    return anonymized.text


def redact_txt(file_path: Path) -> str:
    """Read a plain-text file and return redacted content."""
    raw = file_path.read_text(encoding="utf-8", errors="ignore")
    return redact_text(raw)


def redact_docx(file_path: Path) -> "docx.Document":
    """Open a .docx file, redact each paragraph in-place, return the document."""
    if docx is None:
        raise RuntimeError("python-docx not installed")

    doc = docx.Document(str(file_path))

    for para in doc.paragraphs:
        if para.text.strip():
            redacted = redact_text(para.text)
            # Preserve the first run's formatting; replace its text
            if para.runs:
                para.runs[0].text = redacted
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.text = redacted

    # Also redact text inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        redacted = redact_text(para.text)
                        if para.runs:
                            para.runs[0].text = redacted
                            for run in para.runs[1:]:
                                run.text = ""

    return doc


# ---------------------------------------------------------------------------
# File processor
# ---------------------------------------------------------------------------

def process_file(file_path: Path) -> None:
    suffix = file_path.suffix.lower()
    stem = file_path.stem

    log.info("Processing: %s", file_path.name)

    try:
        if suffix == ".txt":
            redacted_content = redact_txt(file_path)
            out = REDACTED / f"{stem}_redacted.txt"
            out.write_text(redacted_content, encoding="utf-8")
            shutil.move(str(file_path), str(ARCHIVE / file_path.name))
            log.info("  -> redacted copy saved to:  %s", out.name)
            log.info("  -> original archived to:    archive/%s", file_path.name)

        elif suffix == ".docx":
            if docx is None:
                raise RuntimeError("python-docx not installed")
            redacted_doc = redact_docx(file_path)
            out = REDACTED / f"{stem}_redacted.docx"
            redacted_doc.save(str(out))
            shutil.move(str(file_path), str(ARCHIVE / file_path.name))
            log.info("  -> redacted copy saved to:  %s", out.name)
            log.info("  -> original archived to:    archive/%s", file_path.name)

        else:
            log.warning("  Unsupported file type (%s) — moved to error/", suffix)
            shutil.move(str(file_path), str(ERROR / file_path.name))

    except Exception as exc:
        log.error("  Failed to process %s: %s", file_path.name, exc)
        try:
            shutil.move(str(file_path), str(ERROR / file_path.name))
        except Exception:
            pass


# ---------------------------------------------------------------------------
# Watchdog event handler
# ---------------------------------------------------------------------------

class ResumeHandler(FileSystemEventHandler):
    def on_created(self, event):
        if event.is_directory:
            return
        file_path = Path(event.src_path)
        # Give the OS time to finish writing / copying the file
        time.sleep(COPY_SETTLE_SECONDS)
        if file_path.exists():
            process_file(file_path)

    # Also handle files moved/renamed into the folder
    def on_moved(self, event):
        if event.is_directory:
            return
        dest = Path(event.dest_path)
        if dest.parent == INCOMING:
            time.sleep(COPY_SETTLE_SECONDS)
            if dest.exists():
                process_file(dest)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    for folder in [INCOMING, REDACTED, ARCHIVE, ERROR]:
        folder.mkdir(parents=True, exist_ok=True)

    log.info("=" * 60)
    log.info("Resume PII Redactor started")
    log.info("Watching: %s", INCOMING)
    log.info("Supported types: .txt, .docx")
    log.info("Press Ctrl+C to stop.")
    log.info("=" * 60)

    handler = ResumeHandler()
    observer = Observer()
    observer.schedule(handler, str(INCOMING), recursive=False)
    observer.start()

    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        log.info("Stopping…")
        observer.stop()

    observer.join()
    log.info("Redactor stopped.")


if __name__ == "__main__":
    main()
